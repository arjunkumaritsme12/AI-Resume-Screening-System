import re
import docx
import PyPDF2
import io

# ============ TEXT CLEANING ============
def cleanResume(txt):
    """
    Clean and normalize resume text by removing URLs, special characters, and extra whitespace.
    """
    if not txt:
        return ""

    # Remove URLs
    txt = re.sub(r'http\S+|www\S+', ' ', txt)
    
    # Remove special characters but keep spaces
    txt = re.sub(r'[^a-zA-Z0-9\s]', ' ', txt)
    
    # Remove extra whitespace
    txt = re.sub(r'\s+', ' ', txt)
    
    return txt.lower().strip()


def clean_text(text):
    """Alias for cleanResume."""
    return cleanResume(text)


# ============ FILE EXTRACTORS ============
def extract_pdf(file_obj):
    """
    Extract text from PDF file.
    
    Args:
        file_obj: File object or file path string or BytesIO object
    
    Returns:
        Extracted text string
    """
    try:
        if isinstance(file_obj, str):
            with open(file_obj, 'rb') as f:
                file_obj = f
                reader = PyPDF2.PdfReader(file_obj)
        else:
            reader = PyPDF2.PdfReader(file_obj)

        text = ""
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            page_text = page.extract_text()
            if page_text:
                text += page_text + " "

        return text.strip()
    except Exception as e:
        print(f"Error extracting PDF: {str(e)}")
        return ""


def extract_docx(file_obj):
    """
    Extract text from DOCX file.
    
    Args:
        file_obj: File object or file path string
    
    Returns:
        Extracted text string
    """
    try:
        if isinstance(file_obj, str):
            doc = docx.Document(file_obj)
        else:
            doc = docx.Document(file_obj)

        text = " ".join([p.text for p in doc.paragraphs if p.text.strip()])
        return text.strip()
    except Exception as e:
        print(f"Error extracting DOCX: {str(e)}")
        return ""


def extract_txt(file_obj):
    """
    Extract text from TXT file.
    
    Args:
        file_obj: File object or file path string
    
    Returns:
        Extracted text string
    """
    try:
        if isinstance(file_obj, str):
            with open(file_obj, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read().strip()
        else:
            if isinstance(file_obj, bytes):
                return file_obj.decode('utf-8', errors='ignore').strip()
            else:
                return file_obj.read().decode('utf-8', errors='ignore').strip()
    except Exception as e:
        print(f"Error extracting TXT: {str(e)}")
        return ""


def extract_from_file(file_obj, file_type=None):
    """
    Extract text from file based on file type.
    
    Args:
        file_obj: File object, path string, or bytes
        file_type: Optional file type ('pdf', 'docx', 'txt')
    
    Returns:
        Extracted text string
    """
    if file_type == 'pdf' or (isinstance(file_obj, str) and file_obj.lower().endswith('.pdf')):
        return extract_pdf(file_obj)
    elif file_type == 'docx' or (isinstance(file_obj, str) and file_obj.lower().endswith('.docx')):
        return extract_docx(file_obj)
    elif file_type == 'txt' or (isinstance(file_obj, str) and file_obj.lower().endswith('.txt')):
        return extract_txt(file_obj)
    else:
        # Try to detect from content
        if hasattr(file_obj, 'name'):
            name = file_obj.name.lower()
            if name.endswith('.pdf'):
                return extract_pdf(file_obj)
            elif name.endswith('.docx'):
                return extract_docx(file_obj)
            else:
                return extract_txt(file_obj)
        else:
            return extract_txt(file_obj)


# ============ RESUME DETECTION ============
def is_resume(text):
    """
    Check if provided text looks like a resume.
    Returns True if it contains resume-like keywords.
    """
    if not text or len(text) < 200:
        return False

    resume_keywords = [
        "experience", "education", "skills", "projects",
        "objective", "summary", "certification",
        "internship", "work experience", "worked",
        "technical skills", "responsibilities",
        "achievement", "accomplishment",
        "degree", "university", "college",
        "qualification", "language"
    ]

    text_lower = text.lower()
    matched = sum(1 for word in resume_keywords if word in text_lower)

    return matched >= 3


def is_job_description(text):
    """
    Check if provided text looks like a job description.
    """
    if not text or len(text) < 200:
        return False

    jd_keywords = [
        "requirement", "requirement", "responsible", "duty",
        "description", "salary", "benefit", "location",
        "experience required", "qualifications",
        "skill", "must have", "nice to have",
        "apply", "position", "role", "job"
    ]

    text_lower = text.lower()
    matched = sum(1 for word in jd_keywords if word in text_lower)

    return matched >= 2


# ============ TEXT VALIDATION ============
def validate_text(text, min_length=100):
    """
    Validate if text is suitable for processing.
    """
    if not text:
        return False, "Text is empty"
    
    if len(text) < min_length:
        return False, f"Text is too short (minimum {min_length} characters)"
    
    return True, "Valid"


# ============ UTILITY FUNCTIONS ============
def get_file_size_mb(file_obj):
    """Get file size in MB."""
    try:
        if hasattr(file_obj, 'size'):
            return file_obj.size / (1024 * 1024)
        elif hasattr(file_obj, 'tell') and hasattr(file_obj, 'seek'):
            current_pos = file_obj.tell()
            file_obj.seek(0, 2)
            size = file_obj.tell()
            file_obj.seek(current_pos)
            return size / (1024 * 1024)
    except:
        pass
    return 0


def truncate_text(text, max_length=1000):
    """Truncate text to maximum length."""
    if len(text) > max_length:
        return text[:max_length] + "..."
    return text
